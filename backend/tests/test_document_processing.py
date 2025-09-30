"""
Integration tests for document processing endpoints.

Run with: pytest backend/tests/test_document_processing.py -v
"""

import os
import io
import pytest
import requests
from pathlib import Path


BASE_URL = "http://localhost:8000"
TEST_SESSION_PREFIX = "test_doc"


def generate_test_pdf():
    """Generate a simple test PDF in memory."""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        c.drawString(100, 750, "Test Document")
        c.drawString(100, 730, "Name: Juan Pérez")
        c.drawString(100, 710, "Email: juan.perez@email.com")
        c.drawString(100, 690, "Phone: +34 612 345 678")
        c.save()
        buffer.seek(0)
        return buffer.getvalue()
    except ImportError:
        pytest.skip("reportlab not installed")


def generate_test_docx():
    """Generate a simple test Word document in memory."""
    try:
        from docx import Document
        
        buffer = io.BytesIO()
        doc = Document()
        doc.add_paragraph("Test Document")
        doc.add_paragraph("Name: María López")
        doc.add_paragraph("Email: maria.lopez@email.com")
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except ImportError:
        pytest.skip("python-docx not installed")


def generate_test_xlsx():
    """Generate a simple test Excel document in memory."""
    try:
        from openpyxl import Workbook
        
        buffer = io.BytesIO()
        wb = Workbook()
        ws = wb.active
        ws['A1'] = 'Name'
        ws['B1'] = 'Email'
        ws['A2'] = 'Carlos García'
        ws['B2'] = 'carlos.garcia@email.com'
        wb.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except ImportError:
        pytest.skip("openpyxl not installed")


class TestDocumentProcessing:
    """Test suite for document processing endpoints."""
    
    @pytest.fixture(autouse=True)
    def setup(self):
        """Setup for each test."""
        self.session_id = f"{TEST_SESSION_PREFIX}_{os.urandom(4).hex()}"
        yield
        try:
            requests.delete(f"{BASE_URL}/sessions/{self.session_id}")
        except:
            pass
    
    def test_health_check(self):
        """Test that the API is running."""
        response = requests.get(f"{BASE_URL}/health")
        assert response.status_code == 200
        assert response.json()["success"] == True
    
    def test_redis_connection(self):
        """Test that Redis is connected."""
        response = requests.get(f"{BASE_URL}/health/redis")
        assert response.status_code == 200
        data = response.json()
        assert data["connected"] == True
        assert data["redis_status"] == "healthy"
    
    def test_process_pdf(self):
        """Test PDF document processing."""
        pdf_content = generate_test_pdf()
        
        files = {'file': ('test.pdf', pdf_content, 'application/pdf')}
        data = {'session_id': self.session_id}
        
        response = requests.post(
            f"{BASE_URL}/document/process",
            files=files,
            data=data
        )
        
        assert response.status_code == 200
        result = response.json()
        assert result["success"] == True
        assert result["pii_detected"] == True
        assert result["entities_anonymized"] > 0
        assert "anonymized_text" in result
        assert "mapping" in result
        assert isinstance(result["mapping"], dict)
        assert result["document_info"]["detected_type"] == "pdf"
    
    def test_process_word(self):
        """Test Word document processing."""
        docx_content = generate_test_docx()
        
        files = {'file': ('test.docx', docx_content, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        data = {'session_id': self.session_id}
        
        response = requests.post(
            f"{BASE_URL}/document/process",
            files=files,
            data=data
        )
        
        assert response.status_code == 200
        result = response.json()
        assert result["success"] == True
        assert result["document_info"]["detected_type"] == "word"
        assert "mapping" in result
    
    def test_process_excel(self):
        """Test Excel document processing."""
        xlsx_content = generate_test_xlsx()
        
        files = {'file': ('test.xlsx', xlsx_content, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')}
        data = {'session_id': self.session_id}
        
        response = requests.post(
            f"{BASE_URL}/document/process",
            files=files,
            data=data
        )
        
        assert response.status_code == 200
        result = response.json()
        assert result["success"] == True
        assert result["document_info"]["detected_type"] == "excel"
        assert "mapping" in result
    
    def test_auto_generate_session_id(self):
        """Test automatic session ID generation."""
        pdf_content = generate_test_pdf()
        
        files = {'file': ('test.pdf', pdf_content, 'application/pdf')}
        
        response = requests.post(
            f"{BASE_URL}/document/process",
            files=files
        )
        
        assert response.status_code == 200
        result = response.json()
        assert "session_id" in result
        assert result["session_id"].startswith("chat_")
    
    def test_get_mapping(self):
        """Test retrieving anonymization mapping."""
        pdf_content = generate_test_pdf()
        
        files = {'file': ('test.pdf', pdf_content, 'application/pdf')}
        data = {'session_id': self.session_id}
        
        process_response = requests.post(
            f"{BASE_URL}/document/process",
            files=files,
            data=data
        )
        assert process_response.status_code == 200
        
        mapping_response = requests.get(
            f"{BASE_URL}/document/mapping/{self.session_id}"
        )
        
        assert mapping_response.status_code == 200
        result = mapping_response.json()
        assert result["success"] == True
        assert "mapping" in result
        assert result["entities_count"] > 0
    
    def test_deanonymize(self):
        """Test deanonymization of processed document."""
        pdf_content = generate_test_pdf()
        
        files = {'file': ('test.pdf', pdf_content, 'application/pdf')}
        data = {'session_id': self.session_id}
        
        process_response = requests.post(
            f"{BASE_URL}/document/process",
            files=files,
            data=data
        )
        assert process_response.status_code == 200
        anonymized_text = process_response.json()["anonymized_text"]
        
        deanon_data = {
            'session_id': self.session_id,
            'anonymized_text': anonymized_text
        }
        
        deanon_response = requests.post(
            f"{BASE_URL}/document/deanonymize",
            data=deanon_data
        )
        
        assert deanon_response.status_code == 200
        result = deanon_response.json()
        assert result["success"] == True
        assert "original_text" in result
        assert result["entities_restored"] > 0
    
    def test_session_persistence(self):
        """Test that session data persists in Redis."""
        pdf_content = generate_test_pdf()
        
        files = {'file': ('test.pdf', pdf_content, 'application/pdf')}
        data = {'session_id': self.session_id}
        
        requests.post(f"{BASE_URL}/document/process", files=files, data=data)
        
        status_response = requests.get(
            f"{BASE_URL}/sessions/{self.session_id}/status"
        )
        
        assert status_response.status_code == 200
        result = status_response.json()
        assert result["exists"] == True
        assert result["status"] == "active"
        assert result["map_size"] > 0
    
    def test_invalid_file_type(self):
        """Test rejection of unsupported file types."""
        files = {'file': ('test.txt', b'Plain text content', 'text/plain')}
        data = {'session_id': self.session_id}
        
        response = requests.post(
            f"{BASE_URL}/document/process",
            files=files,
            data=data
        )
        
        assert response.status_code == 400
    
    def test_empty_file(self):
        """Test rejection of empty files."""
        files = {'file': ('test.pdf', b'', 'application/pdf')}
        data = {'session_id': self.session_id}
        
        response = requests.post(
            f"{BASE_URL}/document/process",
            files=files,
            data=data
        )
        
        assert response.status_code == 400


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])