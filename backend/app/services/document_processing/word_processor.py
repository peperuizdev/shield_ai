"""
Word Document Processor for Shield AI

Specialized processor for Word (.docx) files using python-docx for text extraction.
Handles paragraphs, tables, headers and footers in proper document order.
"""

import io
import logging
from typing import Dict, List, Any
from zipfile import BadZipFile

try:
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

from .base import DocumentProcessor, DocumentProcessingError, DocumentValidationError

logger = logging.getLogger(__name__)


class WordProcessor(DocumentProcessor):
    """
    Word document processor using python-docx for text extraction.
    
    Implements the DocumentProcessor interface for .docx files,
    providing structured text extraction from paragraphs and tables.
    """
    
    def __init__(self):
        """Initialize Word processor."""
        super().__init__()
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError("python-docx is required for Word processing. Install with: pip install python-docx==1.1.0")
    
    def extract_text(self, file_content: bytes) -> str:
        """
        Extract text from Word document using python-docx.
        
        Args:
            file_content (bytes): Raw .docx file bytes
            
        Returns:
            str: Extracted plain text from document
            
        Raises:
            DocumentProcessingError: If text extraction fails
        """
        try:
            self.logger.info("Extracting text from Word document")
            
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            extracted_text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    extracted_text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                extracted_text += "\n"
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    
                    if row_text:
                        extracted_text += " | ".join(row_text) + "\n"
                extracted_text += "\n"
            
            # Extract headers and footers
            for section in doc.sections:
                # Header
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            extracted_text += paragraph.text + "\n"
                
                # Footer
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            extracted_text += paragraph.text + "\n"
            
            if not extracted_text.strip():
                raise DocumentProcessingError("No text could be extracted from Word document")
            
            self.logger.info("Word text extraction completed successfully")
            return extracted_text.strip()
            
        except Exception as e:
            self.logger.error(f"Word text extraction failed: {str(e)}")
            raise DocumentProcessingError(f"Failed to extract text from Word document: {str(e)}")
    
    def validate_file(self, file_content: bytes, filename: str) -> bool:
        """
        Validate Word file structure and readability.
        
        Args:
            file_content (bytes): Raw file bytes
            filename (str): Original filename
            
        Returns:
            bool: True if file is valid .docx
            
        Raises:
            DocumentValidationError: If validation fails
        """
        try:
            if len(file_content) == 0:
                raise DocumentValidationError("Word file is empty")
            
            # Check if it's a .docx file (ZIP-based format)
            if not file_content.startswith(b'PK'):
                raise DocumentValidationError("File is not a valid .docx format (legacy .doc files not supported)")
            
            docx_file = io.BytesIO(file_content)
            
            try:
                doc = Document(docx_file)
                
                # Test if we can access basic document structure
                paragraph_count = len(doc.paragraphs)
                table_count = len(doc.tables)
                
                self.logger.debug(f"Word document has {paragraph_count} paragraphs and {table_count} tables")
                
                # Ensure document has some content
                has_content = False
                
                # Check paragraphs for content
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        has_content = True
                        break
                
                # Check tables for content if no paragraph content
                if not has_content:
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    has_content = True
                                    break
                            if has_content:
                                break
                        if has_content:
                            break
                
                if not has_content:
                    self.logger.warning(f"Word document {filename} appears to be empty")
            
            except PackageNotFoundError:
                raise DocumentValidationError("Invalid .docx file structure")
            except BadZipFile:
                raise DocumentValidationError("Corrupted .docx file (invalid ZIP structure)")
            
            self.logger.info(f"Word validation successful: {filename}")
            return True
            
        except DocumentValidationError:
            raise
        except Exception as e:
            self.logger.error(f"Word validation failed for {filename}: {str(e)}")
            raise DocumentValidationError(f"Word validation failed: {str(e)}")
    
    def get_supported_extensions(self) -> List[str]:
        """
        Get list of Word file extensions supported.
        
        Returns:
            List[str]: Supported extensions
        """
        return ['.docx', '.DOCX']
    
    def get_supported_mime_types(self) -> List[str]:
        """
        Get list of Word MIME types supported.
        
        Returns:
            List[str]: Supported MIME types
        """
        return [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        ]
    
    def get_metadata(self, file_content: bytes) -> Dict[str, Any]:
        """
        Extract metadata from Word document.
        
        Args:
            file_content (bytes): Raw file bytes
            
        Returns:
            Dict[str, Any]: Word document metadata
        """
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            metadata = {
                "file_size": len(file_content),
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "section_count": len(doc.sections)
            }
            
            # Extract core properties if available
            if hasattr(doc, 'core_properties') and doc.core_properties:
                props = doc.core_properties
                
                if props.title:
                    metadata["title"] = props.title
                if props.author:
                    metadata["author"] = props.author
                if props.subject:
                    metadata["subject"] = props.subject
                if props.created:
                    metadata["created"] = props.created.isoformat() if props.created else None
                if props.modified:
                    metadata["modified"] = props.modified.isoformat() if props.modified else None
                if props.last_modified_by:
                    metadata["last_modified_by"] = props.last_modified_by
            
            # Calculate text statistics
            total_chars = 0
            total_words = 0
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text = paragraph.text.strip()
                    total_chars += len(text)
                    total_words += len(text.split())
            
            metadata.update({
                "total_characters": total_chars,
                "total_words": total_words
            })
            
            return metadata
            
        except Exception as e:
            self.logger.warning(f"Could not extract Word metadata: {str(e)}")
            return {"file_size": len(file_content)}